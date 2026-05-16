"""
Извлечение текста из вложений к заявке на изменение резюме.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from sqlalchemy import text
from sqlalchemy.orm import Session

from app.core.config import get_settings


@dataclass(frozen=True)
class DocumentTextExtractionResult:
    """
    Результат попытки извлечения текста из файла.
    """
    extracted_text: str | None
    extraction_status: str
    message: str
    is_text_extracted: bool


class DocumentTextExtractionService:
    """
    Сервис извлечения текста из документа для semantic routing.
    """

    MAX_TEXT_CHARS = 200_000

    IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}
    DOCX_EXTENSIONS = {".docx"}
    PDF_EXTENSIONS = {".pdf"}
    LEGACY_DOC_EXTENSIONS = {".doc"}

    def __init__(self) -> None:
        self.settings = get_settings()

    def extract_text(
        self,
        file_path: str | Path,
        mime_type: str | None = None,
        original_filename: str | None = None,
    ) -> DocumentTextExtractionResult:
        """
        Извлекает текст из файла.
        """
        if not self.settings.document_extraction_enabled:
            return DocumentTextExtractionResult(
                extracted_text=None,
                extraction_status="failed",
                message="Извлечение текста из документов отключено настройкой DOCUMENT_EXTRACTION_ENABLED.",
                is_text_extracted=False,
            )

        path = Path(file_path)

        if not path.exists() or not path.is_file():
            return DocumentTextExtractionResult(
                extracted_text=None,
                extraction_status="failed",
                message="Файл не найден, текст вложения не извлечён.",
                is_text_extracted=False,
            )

        extension_source = original_filename or path.name
        extension = Path(extension_source).suffix.lower().strip()
        normalized_mime = (mime_type or "").strip().lower()

        try:
            if extension in self.PDF_EXTENSIONS or normalized_mime == "application/pdf":
                extracted_text = self._extract_pdf_text(path)
                return self._build_text_result(extracted_text, "PDF")

            if (
                extension in self.DOCX_EXTENSIONS
                or normalized_mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                extracted_text = self._extract_docx_text(path)
                return self._build_text_result(extracted_text, "DOCX")

            if extension in self.IMAGE_EXTENSIONS or normalized_mime in {"image/jpeg", "image/png"}:
                return DocumentTextExtractionResult(
                    extracted_text=None,
                    extraction_status="failed",
                    message="Текст из изображения не извлечён, используется описание заявки.",
                    is_text_extracted=False,
                )

            if extension in self.LEGACY_DOC_EXTENSIONS or normalized_mime == "application/msword":
                return DocumentTextExtractionResult(
                    extracted_text=None,
                    extraction_status="failed",
                    message="Текст из DOC не извлечён: формат требует внешнего конвертера. Используется описание заявки.",
                    is_text_extracted=False,
                )

            return DocumentTextExtractionResult(
                extracted_text=None,
                extraction_status="failed",
                message="Формат файла не поддерживает извлечение текста, используется описание заявки.",
                is_text_extracted=False,
            )

        except Exception as exc:
            return DocumentTextExtractionResult(
                extracted_text=None,
                extraction_status="failed",
                message=f"Не удалось извлечь текст из вложения: {exc}",
                is_text_extracted=False,
            )

    def extract_and_update_document(
        self,
        session: Session,
        document_id: int,
        file_path: str | Path,
        mime_type: str | None = None,
        original_filename: str | None = None,
    ) -> DocumentTextExtractionResult:
        """
        Извлекает текст и обновляет employee_documents, если нужные поля существуют.
        """
        result = self.extract_text(
            file_path=file_path,
            mime_type=mime_type,
            original_filename=original_filename,
        )
        self.update_document_extraction_result(
            session=session,
            document_id=document_id,
            extraction_result=result,
        )
        return result

    def update_document_extraction_result(
        self,
        session: Session,
        document_id: int,
        extraction_result: DocumentTextExtractionResult,
    ) -> None:
        """
        Обновляет employee_documents.extracted_text / extraction_status, если поля есть.
        """
        columns = self._get_table_columns(session, "employee_documents")

        if "extracted_text" not in columns and "extraction_status" not in columns:
            return

        assignments: list[str] = []
        params = {
            "document_id": int(document_id),
            "extracted_text": extraction_result.extracted_text,
            "extraction_status": self._normalize_extraction_status(extraction_result.extraction_status),
        }

        if "extracted_text" in columns:
            assignments.append("extracted_text = :extracted_text")

        if "extraction_status" in columns:
            assignments.append("extraction_status = :extraction_status")

        if "updated_at" in columns:
            assignments.append("updated_at = NOW()")

        if not assignments:
            return

        query = text(
            f"""
            UPDATE employee_documents
            SET {", ".join(assignments)}
            WHERE id = :document_id
            """
        )
        session.execute(query, params)

    def _extract_pdf_text(self, file_path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("Не установлен пакет pypdf. Добавьте pypdf в requirements.txt.") from exc

        reader = PdfReader(str(file_path))
        page_texts: list[str] = []

        for page in reader.pages:
            page_text = page.extract_text() or ""
            page_text = page_text.strip()
            if page_text:
                page_texts.append(page_text)

        return self._normalize_text("\n\n".join(page_texts))

    def _extract_docx_text(self, file_path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("Не установлен пакет python-docx. Добавьте python-docx в requirements.txt.") from exc

        document = Document(str(file_path))
        parts: list[str] = []

        for paragraph in document.paragraphs:
            text_value = (paragraph.text or "").strip()
            if text_value:
                parts.append(text_value)

        for table in document.tables:
            for row in table.rows:
                row_cells = [
                    (cell.text or "").strip()
                    for cell in row.cells
                    if (cell.text or "").strip()
                ]
                if row_cells:
                    parts.append(" | ".join(row_cells))

        return self._normalize_text("\n".join(parts))

    def _build_text_result(self, extracted_text: str, source_label: str) -> DocumentTextExtractionResult:
        normalized_text = self._normalize_text(extracted_text)

        if not normalized_text:
            return DocumentTextExtractionResult(
                extracted_text=None,
                extraction_status="failed",
                message=f"Текст из {source_label} не найден, используется описание заявки.",
                is_text_extracted=False,
            )

        return DocumentTextExtractionResult(
            extracted_text=normalized_text,
            extraction_status="processed",
            message=f"Текст из {source_label} успешно извлечён.",
            is_text_extracted=True,
        )

    def _normalize_text(self, value: str | None) -> str:
        normalized = "\n".join(
            line.strip()
            for line in (value or "").replace("\r", "\n").split("\n")
            if line.strip()
        )
        return normalized[: self.MAX_TEXT_CHARS]

    def _normalize_extraction_status(self, value: str | None) -> str:
        normalized = (value or "").strip().lower()
        if normalized in {"pending", "processed", "failed"}:
            return normalized
        return "failed"

    def _get_table_columns(self, session: Session, table_name: str) -> set[str]:
        query = text(
            """
            SELECT column_name AS column_name
            FROM information_schema.columns
            WHERE table_schema = DATABASE()
              AND table_name = :table_name
            """
        )

        rows = session.execute(query, {"table_name": table_name}).mappings().all()
        return {str(row["column_name"]) for row in rows if row.get("column_name")}